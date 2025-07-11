from pathlib import Path
from typing import List, Dict, Any, Union, Optional
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph

from file_utilities.core.file import File


class DocxFile(File):
    """
    Generic Docx File class.

    Parameters
    ----------
    path:
        Path to the file.
    """

    def __init__(self, path: Union[str, Path]):
        super().__init__(path)
        self._validate_docx_extension(self.path)
        self._load_docx()

    @property
    def paragraph_count(self) -> int:
        """
        The number of paragraphs in the document. Note that this includes empty
        paragraphs, which may be used for spacing.

        Returns
        -------
        paragraph_count:
            The number of paragraphs in the document.
        """
        return len(self.document.paragraphs)

    @property
    def word_count(self) -> int:
        """
        The number of words in the document. Note that this is a simple count
        based on whitespace separation and may not be accurate for all languages.
        It also includes words in headers, footers, and other non-body text.

        Returns
        -------
        word_count:
            The number of words in the document.
        """
        return sum(len(p.text.split()) for p in self.document.paragraphs)

    def _load_docx(self):
        """Loads (or reloads) the docx file, updating the document attr."""
        self.document = Document(self.path)

    @staticmethod
    def _validate_docx_extension(path: Union[str, Path]):
        """
        Check if the provided file is a .docx file.

        Parameters
        ----------
        path:
            Path object representing the file path to check.
        """

        path = Path(path)

        if path.suffix.lower() != ".docx":
            raise ValueError(f"Provided file must be a .docx file. Got: {path}")

    def load_basic_content(self) -> List[Dict[str, Any]]:
        """
        Loads the contents of a DOCX file preserving only basic formatting like
        bold, italics, alignment, and paragraphs. Returns a list of dictionaries
        representing the content.

        Returns
        -------
        content:
            A list of dictionaries representing the document content with
            formatting.
        """
        content = []

        for paragraph in self.document.paragraphs:
            paragraph_data = self._basic_parse_paragraph(paragraph)
            content.append(paragraph_data)

        return content

    def _basic_parse_paragraph(self, paragraph: Paragraph) -> Dict[str, Any]:
        """
        Parse a paragraph and extract text with formatting.

        Parameters
        ----------
        paragraph:
            A docx Paragraph object.

        Returns
        -------
        paragraph_data:
            A dictionary representing the paragraph with formatting.
        """
        runs_data = []

        for run in paragraph.runs:
            run_data = {
                "text": run.text,
                "bold": run.bold or False,
                "italic": run.italic or False,
            }
            runs_data.append(run_data)

        # Get paragraph alignment
        alignment = self._get_alignment(paragraph.alignment)

        return {
            "type": "paragraph",
            "alignment": alignment,
            "runs": runs_data,
            "text": paragraph.text,
            "is_empty": len(paragraph.text.strip()) == 0,
        }

    def _get_alignment(self, alignment: Optional[WD_PARAGRAPH_ALIGNMENT]) -> str:
        """
        Convert docx alignment to string.

        Parameters
        ----------
        alignment:
            A WD_PARAGRAPH_ALIGNMENT value or None.

        Returns
        -------
        alignment_str:
            A string representing the alignment ("left", "center", "right",
            "justified").
        """
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justified",
            None: "left",
        }
        return alignment_map.get(alignment, "left")

    def write(self, data: Optional[bytes] = None):
        """
        Writes data to the file.

        Parameters
        ----------
        data:
            The data to write to the file. If None, the current file contents
            are saved (updating its metadata).
        """
        raise NotImplementedError("DocxFile is read-only")

    def __repr__(self) -> str:
        return f"DocxFile(path={self.path}, paragraphs={self.paragraph_count})"
